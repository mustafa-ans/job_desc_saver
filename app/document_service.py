from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def format_document(doc, job_text, job_title, company_name, company_website=""):
    title = doc.add_heading(job_title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    metadata = doc.add_paragraph()
    metadata.add_run(f"Company: {company_name}").bold = True

    if company_website:
        metadata.add_run("\n")
        run = metadata.add_run(f"Website: {company_website}")
        run.font.color.rgb = RGBColor(0, 102, 204)

    metadata.add_run(f"\nSaved: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True

    doc.add_paragraph("_" * 80)

    paragraphs = job_text.strip().split('\n\n')

    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        is_heading = False
        heading_keywords = [
            'DEINE ROLLE', 'DEIN PROFIL', 'ABOUT THE JOB', 'YOUR ROLE',
            'YOUR PROFILE', 'REQUIREMENTS', 'RESPONSIBILITIES', 'QUALIFICATIONS',
            'BENEFITS', 'ABOUT US', 'WIR BIETEN', 'AUFGABEN', 'ANFORDERUNGEN',
            'WAS WIR BIETEN', 'DAS BIETEN WIR', 'UNSER VERSPRECHEN'
        ]

        for keyword in heading_keywords:
            if para_text.upper().startswith(keyword):
                is_heading = True
                break

        if is_heading:
            heading = doc.add_heading(para_text, level=2)
            heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        else:
            p = doc.add_paragraph(para_text)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    return doc


def save_to_word(job_text, filename, job_title, company_name, company_website=""):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    format_document(doc, job_text, job_title, company_name, company_website)
    doc.save(filename)
    return filename


def save_to_pdf(docx_filename, pdf_filename):
    try:
        from docx2pdf import convert
        convert(docx_filename, pdf_filename)
        print(f"✓ PDF saved: {pdf_filename}")
        return True
    except ImportError:
        print("⚠ docx2pdf not available. Install with: pip install docx2pdf")
        print("  Note: Requires Microsoft Word on Windows/Mac")
        return False
    except Exception as e:
        print(f"⚠ PDF conversion failed: {e}")
        print("  You can manually open the temp Word file and save as PDF")
        return False


def save_raw_text(job_text, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(job_text)
    print(f"✓ Raw text archived: {filename}")
