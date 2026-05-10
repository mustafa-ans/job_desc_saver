#!/usr/bin/env python3
"""
Job Description Saver
Saves pasted job descriptions to PDF and archives raw text
Copies cover letter into the same job-specific subfolder
Updates local Excel tracker

Supports:
  :back    -> go to previous question
  :restart -> restart metadata entry
  :quit    -> exit program
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import sys
import shutil
import openpyxl
import pyperclip

OUTPUT_DIR = r"REMOVED_LOCAL_PATH"
EXCEL_PATH = r"REMOVED_LOCAL_PATH"


def clear_console():
    """Clear terminal screen once at script start"""
    os.system('cls' if os.name == 'nt' else 'clear')


def format_document(doc, job_text, job_title, company_name, company_website=""):
    """Add formatted content to the Word document"""

    title = doc.add_heading(job_title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    metadata = doc.add_paragraph()
    metadata.add_run(f"Company: {company_name}").bold = True

    if company_website:
        metadata.add_run("\n")
        run = metadata.add_run(f"Website: {company_website}")
        run.font.color.rgb = RGBColor(0, 102, 204)

    metadata.add_run(f"\nSaved: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True

    doc.add_paragraph("_" * 80)

    paragraphs = job_text.strip().split('\n\n')

    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        is_heading = False
        heading_keywords = [
            'DEINE ROLLE', 'DEIN PROFIL', 'ABOUT THE JOB', 'YOUR ROLE',
            'YOUR PROFILE', 'REQUIREMENTS', 'RESPONSIBILITIES', 'QUALIFICATIONS',
            'BENEFITS', 'ABOUT US', 'WIR BIETEN', 'AUFGABEN', 'ANFORDERUNGEN',
            'WAS WIR BIETEN', 'DAS BIETEN WIR', 'UNSER VERSPRECHEN'
        ]

        for keyword in heading_keywords:
            if para_text.upper().startswith(keyword):
                is_heading = True
                break

        if is_heading:
            heading = doc.add_heading(para_text, level=2)
            heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        else:
            p = doc.add_paragraph(para_text)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    return doc


def save_to_word(job_text, filename, job_title, company_name, company_website=""):
    """Save job description to Word document"""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    format_document(doc, job_text, job_title, company_name, company_website)
    doc.save(filename)
    return filename


def save_to_pdf(docx_filename, pdf_filename):
    """Convert Word document to PDF"""
    try:
        from docx2pdf import convert
        convert(docx_filename, pdf_filename)
        print(f"✓ PDF saved: {pdf_filename}")
        return True
    except ImportError:
        print("⚠ docx2pdf not available. Install with: pip install docx2pdf")
        print("  Note: Requires Microsoft Word on Windows/Mac")
        return False
    except Exception as e:
        print(f"⚠ PDF conversion failed: {e}")
        print("  You can manually open the temp Word file and save as PDF")
        return False


def save_raw_text(job_text, filename):
    """Save raw job description text as-is to archive"""
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(job_text)
    print(f"✓ Raw text archived: {filename}")


def copy_to_clipboard(text):
    """Copy plain text to clipboard"""
    try:
        pyperclip.copy(text)
        print(f"✓ Copied to clipboard: {text}")
    except Exception as e:
        print(f"⚠ Could not copy to clipboard: {e}")


def get_portal_name(choice):
    mapping = {
        "1": "Linkedin",
        "2": "Direct",
        "3": "Website",
        "4": "Email",
        "5": "Linkedin Easy Apply"
    }
    return mapping.get(choice, "Unknown")


def append_to_excel(row_data):
    """Appends data to the local Excel tracker"""
    try:
        wb = openpyxl.load_workbook(EXCEL_PATH)
        ws = wb.active
        ws.append(row_data)
        wb.save(EXCEL_PATH)
        return True
    except Exception as e:
        print(f"✗ Error updating Excel: {e}")
        return False


def move_cover_letter_to_folder(source_path, destination_folder):
    """Move cover letter file into the job folder"""
    try:
        if not source_path.strip():
            print("⚠ No cover letter path provided. Skipping move.")
            return ""

        source_path = source_path.strip().strip('"')

        if not os.path.isfile(source_path):
            print(f"⚠ Cover letter file not found: {source_path}")
            return ""

        destination_path = os.path.join(destination_folder, os.path.basename(source_path))
        moved_path = shutil.move(source_path, destination_path)
        print(f"✓ Cover letter moved to: {moved_path}")
        return moved_path
    except Exception as e:
        print(f"⚠ Error moving cover letter: {e}")
        return ""


def delete_exact_matching_docx_for_pdf(pdf_path):
    """
    Delete only the DOCX file with the exact same basename as the PDF,
    in the same folder.
    Example:
      Acme_Backend_Developer.pdf -> delete Acme_Backend_Developer.docx
    """
    try:
        pdf_root, pdf_ext = os.path.splitext(pdf_path)
        if pdf_ext.lower() != ".pdf":
            return False

        matching_docx = pdf_root + ".docx"

        if os.path.isfile(matching_docx):
            os.remove(matching_docx)
            print(f"✓ Deleted exact matching Word file: {matching_docx}")
            return True
        else:
            # if not found then delete this file "REMOVED_LOCAL_PATH"
            cover_letter_docx = os.path.join(os.path.dirname(pdf_path), "Cover_Letter_Mustafa_Ansari.docx")
            if os.path.isfile(cover_letter_docx):
                os.remove(cover_letter_docx)
                print(f"✓ Deleted cover letter Word file: {cover_letter_docx}")
                return True
            return False
    except Exception as e:
        print(f"⚠ Error deleting exact matching Word file: {e}")
        return False


def collect_job_metadata():
    """
    Collect company name, website, and job title with support for:
    :back    -> go to previous question
    :restart -> restart metadata entry
    :quit    -> exit program
    """
    fields = [
        ("company_name", "Enter company name (required): ", True),
        ("company_website", "Enter company website (optional, press Enter to skip): ", False),
        ("job_title", "Enter job title (required): ", True),
    ]

    answers = {
        "company_name": "",
        "company_website": "",
        "job_title": "",
    }

    index = 0

    while index < len(fields):
        field_key, prompt, required = fields[index]
        current_value = answers[field_key]

        if current_value:
            user_input = input(f"{prompt}[current: {current_value}] ").strip()
        else:
            user_input = input(prompt).strip()

        cmd = user_input.lower()

        if cmd == ":quit":
            print("Exiting...")
            sys.exit(0)

        if cmd == ":restart":
            print("Restarting metadata entry...\n")
            answers = {k: "" for k in answers}
            index = 0
            continue

        if cmd == ":cls":
            clear_console()
            print("Metadata entry restarted (console cleared).\n")
            answers = {k: "" for k in answers}
            index = 0
            continue

        if cmd == ":back":
            if index == 0:
                print("Already at the first question.")
            else:
                index -= 1
                print("Going back to previous question.")
            continue

        if required and not user_input:
            print("This field is required!")
            continue

        answers[field_key] = user_input
        index += 1

    return answers["company_name"], answers["company_website"], answers["job_title"]


def confirm_metadata(company_name, company_website, job_title):
    """Show entered metadata and ask for confirmation"""
    while True:
        print("\nPlease confirm:")
        print(f"Company name   : {company_name}")
        print(f"Company website: {company_website or '(none)'}")
        print(f"Job title      : {job_title}")

        confirm = input("Is this correct? (y/n): ").strip().lower()
        if confirm == "y":
            return True
        elif confirm == "n":
            return False
        else:
            return True

def sanitize_name(text):
    text = "".join(c for c in text if c.isalnum() or c in (' ', '-', '_')).strip()
    text = "_".join(text.split())
    return text or "untitled"


def build_safe_base_name(output_dir, company_name, job_title, extra_suffix=""):
    """
    Build a base filename that keeps the final full path under Windows MAX_PATH.
    extra_suffix can be used for things like '.pdf', '.txt', or '\\temp_' planning.
    """
    MAX_FULL_PATH = 259

    safe_company = sanitize_name(company_name)
    safe_job_title = sanitize_name(job_title)

    separator = "_"
    candidate = f"{safe_company}{separator}{safe_job_title}"

    # Reserve length for directory + slash + filename + optional suffix
    reserved_len = len(os.path.join(output_dir, "")) + len(extra_suffix)
    max_base_len = MAX_FULL_PATH - reserved_len

    if max_base_len < 20:
        raise ValueError("Output directory path is too long to safely create files.")

    if len(candidate) <= max_base_len:
        return candidate

    # Split available space roughly across both parts
    # account for the underscore between them
    available = max_base_len - len(separator)

    # Ensure both sides keep some readable content
    min_each = 8
    company_share = max(min_each, available // 2)
    title_share = max(min_each, available - company_share)

    if company_share + title_share > available:
        title_share = available - company_share

    safe_company = safe_company[:company_share].rstrip("_- ")
    safe_job_title = safe_job_title[:title_share].rstrip("_- ")

    candidate = f"{safe_company}{separator}{safe_job_title}"

    if len(candidate) > max_base_len:
        candidate = candidate[:max_base_len].rstrip("_- ")

    return candidate

def main():
    clear_console()

    print("=" * 80)
    print("JOB DESCRIPTION SAVER")
    print("=" * 80)
    print()

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    print("Commands available while answering metadata:")
    print("  :back    -> redo previous answer")
    print("  :restart -> restart all metadata fields")
    print("  :cls     -> clear console")
    print("  :quit    -> exit program")
    print()

    while True:
        company_name, company_website, job_title = collect_job_metadata()
        if confirm_metadata(company_name, company_website, job_title):
            break
        print("\nLet's enter them again.\n")

    print()
    print("Paste the job description below.")
    print("When finished, type 'END' on a new line and press Enter")
    print("-" * 80)

    try:
        lines = []
        while True:
            try:
                line = input()
                if line.strip().upper() == 'END':
                    break
                lines.append(line)
            except EOFError:
                break
        job_text = '\n'.join(lines)
    except KeyboardInterrupt:
        print("\n\nOperation cancelled.")
        sys.exit(0)

    if not job_text.strip():
        print("\nError: No text provided!")
        sys.exit(1)

    print()
    print("-" * 80)

    base_filename = build_safe_base_name(OUTPUT_DIR, company_name, job_title)

    job_folder = os.path.join(OUTPUT_DIR, base_filename)
    os.makedirs(job_folder, exist_ok=True)

    txt_filename = os.path.join(job_folder, f"{base_filename}.txt")
    try:
        save_raw_text(job_text, txt_filename)
    except Exception as e:
        print(f"✗ Error saving text archive: {e}")

    docx_filename = os.path.join(job_folder, f"temp_{base_filename}.docx")
    pdf_filename = os.path.join(job_folder, f"{base_filename}.pdf")

    try:
        save_to_word(job_text, docx_filename, job_title, company_name, company_website)
    except Exception as e:
        print(f"✗ Error creating document: {e}")
        sys.exit(1)

    pdf_success = save_to_pdf(docx_filename, pdf_filename)

    final_created_file = ""

    if pdf_success:
        final_created_file = pdf_filename
        try:
            os.remove(docx_filename)
        except Exception:
            pass
    else:
        final_created_file = docx_filename
        print(f"⚠ Temporary Word file kept: {docx_filename}")
        print("  You can manually convert it to PDF")

    print("\n--- Cover Letter Move ---")
    cover_letter_path = input("Enter full path to cover letter file (optional, press Enter to skip): ").strip()
    moved_cover_letter_path = move_cover_letter_to_folder(cover_letter_path, job_folder)

    if moved_cover_letter_path:
        delete_exact_matching_docx_for_pdf(final_created_file)

    print("\n--- Update Excel Tracker ---")
    job_id = input("Job ID (optional): ").strip()
    print("Portal: 1-Linkedin, 2-Direct, 3-Website, 4-Email")
    portal_choice = input("Select portal (1-4): ").strip()
    job_link = input("Job Link: ").strip()
    job_type = input("Type (Startup/Mid/Big): ").strip()
    industry = input("Industry: ").strip()
    location = input("Location: ").strip()
    status = input("Status: ").strip()
    docs_sent = input("Docs Sent: ").strip()
    todo_done = input("TO-DO (optional): ").strip()
    hr_name = input("HR Name: ").strip()
    hr_contact = input("HR Contact: ").strip()
    comments = input("Comments (optional): ").strip()
    id_pw = input("ID/PW (optional): ").strip()

    row_data = [
        job_title,                              # A: Position
        job_id,                                 # B: ID
        datetime.now().strftime("%Y-%m-%d"),    # C: Date
        "",                                     # D: Interview on
        "Yes",                                  # E: Applied
        get_portal_name(portal_choice),         # F: Portal
        job_link,                               # G: Link
        company_name,                           # H: Company
        job_type,                               # I: Type
        industry,                               # J: Industry
        location,                               # K: Location
        status,                                 # L: Status
        docs_sent,                              # M: Docs
        todo_done,                              # N: TODO
        hr_name,                                # O: HR Name
        hr_contact,                             # P: HR Contact
        comments,                               # Q: Comments
        id_pw,                                  # R: ID/PW
        job_folder                              # S: Path to JD in explorer
    ]

    if append_to_excel(row_data):
        print("✓ Excel tracker updated successfully.")

    print()
    print("=" * 80)
    print("DONE!")
    print(f"Job folder: {job_folder}")
    print("")
    print(f"Main output file: {final_created_file}")
    print("")
    if moved_cover_letter_path:
        print(f"Moved cover letter: {moved_cover_letter_path}")
    print("")
    print(f"Text archive file: {txt_filename}")
    print("=" * 80)


if __name__ == "__main__":
    main()